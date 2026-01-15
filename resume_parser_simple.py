import PyPDF2
from docx import Document
import re
import io

def extract_text_from_pdf(file):
    """
    Extract text from PDF file or file object.
    Preserves structural layout better than standard extraction.
    """
    try:
        # Check if file is a path or file-like object
        if hasattr(file, 'read'):
            reader = PyPDF2.PdfReader(file)
        else:
            with open(file, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                
        text = []
        for page in reader.pages:
            # extract_text() usually preserves newlines, which is good
            page_content = page.extract_text()
            if page_content:
                text.append(page_content)
        
        full_text = "\n".join(text)
        return clean_text(full_text)
        
    except Exception as e:
        print(f"PDF error: {e}")
        return ""

def extract_text_from_docx(file):
    """
    Extract text from DOCX file or file object.
    Iterates through tables and paragraphs to capture all data.
    """
    try:
        doc = Document(file)
        full_text = []
        
        # 1. Extract from standard paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
                
        # 2. Extract from tables (common in creative resumes)
        for table in doc.tables:
            for row in table.rows:
                # Join cells with a pipe | to keep structure
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))

        return clean_text("\n".join(full_text))
        
    except Exception as e:
        print(f"DOCX error: {e}")
        return ""

def clean_text(text):
    """
    Clean text while PRESERVING critical structure (newlines).
    Crucial for AI to understand sections like 'Experience' vs 'Education'.
    """
    if not text:
        return ""
    
    # 1. Normalize unicode characters
    # Replace non-breaking spaces (\xa0) with normal spaces
    text = text.replace('\xa0', ' ')
    # Normalize bullet points to asterisks for consistency
    text = text.replace('\u2022', '*').replace('●', '*').replace('▪', '*')
    
    # 2. Fix multiple spaces but KEEP NEWLINES
    # This regex replaces 2+ spaces/tabs with 1 space, but ignores newlines
    text = re.sub(r'[ \t]+', ' ', text)
    
    # 3. Fix multiple newlines (e.g., 5 enters -> 2 enters)
    # We want to keep paragraph breaks but remove massive empty gaps
    text = re.sub(r'\n\s*\n', '\n\n', text)
    
    # 4. Remove strange control characters but keep printable ones + formatting
    # (Keeps newlines \n and carriage returns \r)
    text = "".join(ch for ch in text if ch.isprintable() or ch in ['\n', '\r'])
    
    return text.strip()